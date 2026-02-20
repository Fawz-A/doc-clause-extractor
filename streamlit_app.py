import streamlit as st
import pdfplumber
from PIL import Image
import pytesseract
import pandas as pd
import re
import tempfile
from pathlib import Path

from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font

st.set_page_config(page_title="Document Clause Converter")

st.title("📄 PDF / Image to Structured Excel Converter")







from docx import Document
from io import BytesIO
from pdf2image import convert_from_path
# -----------------------------
# OCR FOR SCANNED PDFs
# -----------------------------
def extract_text_from_scanned_pdf(pdf_path):
    images = convert_from_path(pdf_path)
    full_text = []

    for img in images:
        text = pytesseract.image_to_string(img, config="--psm 6")
        full_text.append(text)

    return full_text


# -----------------------------
# STRUCTURED WORD DOCUMENT CREATION
# -----------------------------
def create_structured_word(text_pages):
    doc = Document()

    clause_pattern = re.compile(r"^\s*(\d+(?:\.\d+)*\.?)\s+(.*)")

    for page in text_pages:
        lines = page.splitlines()

        for line in lines:
            line = line.strip()
            if not line:
                continue

            match = clause_pattern.match(line)

            if match:
                clause_number = match.group(1)
                content = match.group(2)

                level = clause_number.count(".") + 1
                level = min(level, 4)

                doc.add_heading(f"{clause_number} {content}", level=level)
            else:
                doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer

# -----------------------------
# TEXT EXTRACTION
# -----------------------------







def extract_text_from_pdf(pdf_path: Path):
    pages_text = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text_content = page.extract_text() or ""
            pages_text.append(text_content.strip())
    return pages_text


def extract_text_from_image(image_path: Path):
    img = Image.open(image_path)
    img = img.convert("L")
    text_content = pytesseract.image_to_string(img, config="--psm 6")
    return [text_content.strip()]


# -----------------------------
# CLAUSE PROCESSING
# -----------------------------

def process_text_pages(text_pages):
    rows = []

    clause_pattern = re.compile(r"^\s*(\d+(?:\.\d+)*\.?)\s+(.*)")

    current_clause = None
    current_text = ""

    for page_content in text_pages:
        lines = page_content.splitlines()

        for line in lines:
            line = line.strip()
            if not line:
                continue

            match = clause_pattern.match(line)

            if match:
                if current_clause:
                    rows.append({
                        "clause_number": current_clause,
                        "content": current_text.strip(),
                    })

                current_clause = match.group(1)
                current_text = match.group(2)

            else:
                if current_clause:
                    current_text += " " + line

    if current_clause:
        rows.append({
            "clause_number": current_clause,
            "content": current_text.strip(),
        })

    df = pd.DataFrame(rows)

    # ✅ Add description column
    df["description"] = ""

    return df


# -----------------------------
# STREAMLIT UI
# -----------------------------

# -----------------------------
# STREAMLIT UI
# -----------------------------

tab1, tab2 = st.tabs(["Clause → Excel", "Document → Word"])

# =========================================
# TAB 1 — CLAUSE TO EXCEL (Your Existing Logic)
# =========================================

with tab1:

    uploaded_file = st.file_uploader(
        "Upload PDF or Image",
        type=["pdf", "png", "jpg", "jpeg"],
        key="excel_upload"
    )

    if uploaded_file:

        with st.spinner("Processing document..."):

            suffix = Path(uploaded_file.name).suffix

            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
                tmp_file.write(uploaded_file.read())
                tmp_path = Path(tmp_file.name)

            ext = tmp_path.suffix.lower()

            if ext == ".pdf":
                text_pages = extract_text_from_pdf(tmp_path)
            elif ext in [".png", ".jpg", ".jpeg"]:
                text_pages = extract_text_from_image(tmp_path)
            else:
                st.error("Unsupported file type")
                st.stop()

            df = process_text_pages(text_pages)

        if df.empty:
            st.warning("No numbered clauses detected.")
        else:
            st.success(f"Extracted {len(df)} clauses.")
            st.dataframe(df)

            original_name = Path(uploaded_file.name).stem
            output_filename = f"converted_{original_name}.xlsx"

            output_file = tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False)
            df.to_excel(output_file.name, index=False)

            wb = load_workbook(output_file.name)
            ws = wb.active

            ws.freeze_panes = "A2"

            for cell in ws[1]:
                cell.font = Font(bold=True)

            for row in ws.iter_rows():
                for cell in row:
                    cell.alignment = Alignment(wrap_text=True, vertical="top")

            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter

                for cell in column:
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))

                ws.column_dimensions[column_letter].width = min(max_length + 2, 60)

            wb.save(output_file.name)

            with open(output_file.name, "rb") as f:
                st.download_button(
                    label="⬇ Download Excel File",
                    data=f,
                    file_name=output_filename,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )


# =========================================
# TAB 2 — DOCUMENT TO STRUCTURED WORD
# =========================================

with tab2:

    uploaded_word_file = st.file_uploader(
        "Upload PDF or Image for Word conversion",
        type=["pdf", "png", "jpg", "jpeg"],
        key="word_upload"
    )

    if uploaded_word_file:

        suffix = Path(uploaded_word_file.name).suffix

        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
            tmp_file.write(uploaded_word_file.read())
            tmp_path = Path(tmp_file.name)

        if suffix.lower() == ".pdf":
            text_pages = extract_text_from_pdf(tmp_path)

            # Fallback for scanned PDF
            if not any(text_pages):
                text_pages = extract_text_from_scanned_pdf(tmp_path)
        else:
            text_pages = extract_text_from_image(tmp_path)

        word_buffer = create_structured_word(text_pages)

        st.download_button(
            label="⬇ Download Structured Word File",
            data=word_buffer,
            file_name=f"converted_{Path(uploaded_word_file.name).stem}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )